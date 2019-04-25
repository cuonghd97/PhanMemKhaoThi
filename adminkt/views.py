from django.shortcuts import render
from django.http import HttpResponse, HttpRequest, HttpResponseRedirect,JsonResponse
from django.views import View
from django.contrib.auth import authenticate, login, decorators, logout, get_user
from django.shortcuts import redirect
from django.conf import settings
import json
import os
from wsgiref.util import FileWrapper
from datetime import datetime,date
from docx import Document
from docx.shared import Inches,Pt,Mm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from CoiThi.models import *
# Create your views here.
def home(request):
    user = request.user
    content = {'username': user.tenCanBo}
    if not request.user.is_authenticated:
        return redirect('CoiThi:login')
    return render(request, 'adminkt/index.html',content)
def profile(request):
    user = request.user
    
    data = {'id': request.user.id}
    data.update({'username': request.user.tenCanBo})
    data.update({'tencanbo': request.user.tenCanBo})
    if not request.user.is_authenticated:
        return redirect('CoiThi:login')
    return render(request, 'adminkt/profile.html',data)
def updateprofile(request):
    matkhaumoi = request.POST['matkhaumoi']
    mknhaplai = request.POST['nhaplaimatkhau']
    username = request.user.username
    idcanbo = request.POST['idcanbo']
    if request.POST['nhaplaimatkhau'] != 'notnull':
        request.user.set_password(matkhaumoi)
        request.user.tenCanBo = request.POST['ten']

        request.user.save()

    request.user.tenCanBo = request.POST['ten']

    request.user.save()
    CanBo = authenticate(username = username, password = matkhaumoi)
    login(request, CanBo)
    return HttpResponse('done')
def manage_class(request):
    user = request.user
    if user.is_authenticated and user.position == 0:
        if request.method == 'POST':
            if 'delete' in request.POST:
                LopHoc.objects.get(id=request.POST['delete']).delete()
            elif 'ten' in request.POST:
                # nien_khoa, nam = request.POST['nien_khoa'].split(" - ")
                if request.POST['kieu'] == 'new':
                    try:
                    # print(request.POST['ten'])
                        LopHoc.objects.create(tenLop=request.POST['ten'],maKhoa=KhoaHoc.objects.get(id =request.POST['khoa']),maMon = Mon.objects.get(id = request.POST['mon']))
                    except:
                        pass
                else:
                    try:
                        l = LopHoc.objects.get(id=request.POST['id'])
                        l.tenLop = request.POST['ten']
                        l.maKhoa = KhoaHoc.objects.get(id=request.POST['khoa'])
                        l.maMon = Mon.objects.get(id = request.POST['mon'])
                        l.save()
                    except:
                        pass
        content = {'username': user.tenCanBo, 'ds_khoa': KhoaHoc.objects.all(), 'mon' : Mon.objects.all(),'ds_sinhvien':SinhVien.objects.all(),'ds_lop':LopHoc.objects.all()}
        return render(request, 'adminkt/manager_lop.html', content)
    else:
        return HttpResponseRedirect('/')

def manage_class_data(request):
    user = request.user
    if user.is_authenticated and user.position == 0:
        data = []
        for lop in LopHoc.objects.all():
            ten = '<p id="ten_{}">{}</p>'.format(lop.id, lop.tenLop)
            khoa = '<p id="khoa_{}">{} - {}</p>'.format(lop.id, lop.maKhoa.tenKhoaHoc,lop.maKhoa.he)
            mon = '<p id="mon_{}">{}</p>'.format(lop.id, lop.maMon.tenMon)
            hs = '''
            {1}  <i class="fa fa-list" data-title="{0}" data-toggle="modal" data-target="#detail_student" title = "Danh sách lớp"></i>
            '''.format(lop.tenLop,ChiTietLop.objects.filter(maLop_id = lop.id).count())
            options = '''
                <div class="btn-group">
                    <button type="button" class="btn btn-info" data-toggle="modal" title="Chỉnh sửa" data-target="#new_class" data-title="edit" id="edit_{0}">
                        <i class="fa fa-cog" data-toggle="tooltip" title="Chỉnh sửa"></i>
                    </button>
                    <button type="button" class="btn btn-danger" title="Xóa" data-title="del" id="del_{0}">
                        <i class="fa fa-trash" data-toggle="tooltip" title="Xóa"></i>
                    </button>
                </div>
            '''.format(lop.id)
            data.append([ten,mon, khoa,hs, options])
        big_data = {"data": data}
        json_data = json.loads(json.dumps(big_data))
        return JsonResponse(json_data)
def manage_add_student(request):
    user = request.user
    if user.is_authenticated and user.position == 0:
        if request.method == 'POST':
            if 'delete' in request.POST:
                ChiTietLop.objects.get(id=request.POST['delete']).delete()
            elif 'lop' in request.POST:

                if request.POST['kieu'] == 'new':
                    try:
                        ChiTietLop.objects.create(maLop = LopHoc.objects.get(id = request.POST['lop']) ,maSinhVien=SinhVien.objects.get(id =request.POST['sinhvien']))
                    except:
                        pass
        return JsonResponse({'status':'succsec'})
    else:
        return HttpResponseRedirect('/')
def manage_student_data(request,lop):
    user = request.user
    if user.is_authenticated and user.position == 0:
        data = []
        try:
            ls_chi_tiet = ChiTietLop.objects.filter(maLop = LopHoc.objects.get(tenLop=lop)).values('maLop')
            ls_student = ChiTietLop.objects.filter(maLop__in = ls_chi_tiet)
            now = datetime.datetime.now()
            yearnow = now.year 
            lop = LopHoc.objects.get(tenLop=lop).id
            for student in ls_student:
                fullname = '<p id="full_{0}">{1}</p>'.format(student.maSinhVien.id, student.maSinhVien.tenSinhVien)
                username = '<p id="user_{0}">{1}</p>'.format(student.maSinhVien.id, student.maSinhVien.maSinhVien)
                tuoisv = yearnow - student.maSinhVien.ngaySinh.year
                tuoi = '<p id="tuoi_{}">{}</p>'.format(student.maSinhVien.id,tuoisv)
                capham = '<p id="capham_{0}">{1}</p>'.format(student.maSinhVien.id, student.maSinhVien.maDonVi.tenDonVi)
                diem = '<p id="diem_{0}">{1}</p>'.format(student.maSinhVien.id, student.diem)
                option = '''
                    <div class="btn-group">
                        <button type="button" class="btn btn-info log" data-toggle="modal" title="Lịch sử chỉnh sửa" data-target="#log_student" data-lop="{1}" data-hocvien="{0}" data-title="log" id="log_{0}">
                            <i class="fa fa-history" data-toggle="tooltip" title="Lịch sử chỉnh sửa"></i>
                        </button>
                        <button type="button" class="btn btn-danger" title="Xóa" data-title="del" id="del_{0}">
                            <i class="fa fa-trash" data-toggle="tooltip" title="Xóa"></i>
                        </button>
                    </div>
                '''.format(student.maSinhVien.id,lop)
                data.append([fullname, username,tuoi,capham,diem,option])
        except:
            pass
        big_data = {"data": data}
        json_data = json.loads(json.dumps(big_data))
        return JsonResponse(json_data)

def manage_donvi(request):
    user = request.user
    if user.is_authenticated and user.position == 0:
        if request.method == 'POST':
            if 'delete' in request.POST:
                DonVi.objects.get(id=request.POST['delete']).delete()
            elif 'ten' in request.POST:
                # nien_khoa, nam = request.POST['nien_khoa'].split(" - ")
                if request.POST['kieu'] == 'new':
                    try:
                        DonVi.objects.create(tenDonVi=request.POST['ten'],maDonVi=request.POST['ma'])
                    except:
                        pass
                else:
                    try:
                        l = DonVi.objects.get(id=request.POST['id'])
                        l.tenDonVi = request.POST['ten']
                        l.maDonVi = request.POST['ma']
                        l.save()
                    except:
                        pass
        content = {'username': user.tenCanBo}
        return render(request, 'adminkt/manager_donvi.html', content)
    else:
        return HttpResponseRedirect('/')

def manage_donvi_data(request):
    user = request.user
    if user.is_authenticated and user.position == 0:
        data = []
        for dv in DonVi.objects.all():
            ten = '<p id="ten_{}">{}</p>'.format(dv.id, dv.tenDonVi)
            Madv = '<p id="ma_{}">{}</p>'.format(dv.id, dv.maDonVi)
            options = '''
                <div class="btn-group">
                    <button type="button" class="btn btn-info" title="Chỉnh sửa" data-toggle="modal" data-target="#new_class" data-title="edit" id="edit_{0}">
                        <i class="fa fa-cog" data-toggle="tooltip" title="Chỉnh sửa"></i>
                    </button>
                    <button type="button" class="btn btn-danger" title="Xóa" data-title="del" id="del_{0}">
                        <i class="fa fa-trash" data-toggle="tooltip" title="Xóa"></i>
                    </button>
                </div>
            '''.format(dv.id)
            data.append([ten,Madv,options])
        big_data = {"data": data}
        json_data = json.loads(json.dumps(big_data))
        return JsonResponse(json_data)

def manage_khoa(request):
    user = request.user
    content = {'username': user.tenCanBo}
    if request.user.is_authenticated:
        if request.method == 'POST':
            if 'delete' in request.POST:
                KhoaHoc.objects.get(id=request.POST['delete']).delete()
            else:
                if request.POST['kieu'] == 'new':
                    try:
                        KhoaHoc.objects.create(tenKhoaHoc=request.POST['khoa'], he=request.POST['he'])

                    except:
                        pass
                else:
                    try:
                        nk = KhoaHoc.objects.get(id=request.POST['id'])
                        nk.tenKhoaHoc = request.POST['khoa']
                        nk.he = request.POST['he']
                        nk.save()
                    except:
                        pass
        return render(request,'adminkt/manager_khoa.html',content)
    else:
        return HttpResponseRedirect('/')

def manage_khoa_data(request):
    user = request.user
    if user.is_authenticated and user.position == 0:
        data = []
        for khoa in KhoaHoc.objects.all():
            ten = '<p id="ten_{}">{}</p>'.format(khoa.id, khoa.tenKhoaHoc)
            he = '<p id="he_{}">{}</p>'.format(khoa.id, khoa.he)
            options = '''
                <div class="btn-group">
                    <button type="button" class="btn btn-info" data-toggle="modal" data-target="#new_khoa" data-title="edit" id="edit_{0}">
                        <i class="fa fa-cog" data-toggle="tooltip" title="Chỉnh sửa"></i>
                    </button>
                    <button type="button" class="btn btn-danger" data-title="del" id="del_{0}">
                        <i class="fa fa-trash" data-toggle="tooltip" title="Xóa"></i>
                    </button>
                </div>
            '''.format(khoa.id)
            data.append([ten, he, options])
        big_data = {"data": data}
        json_data = json.loads(json.dumps(big_data))
        return JsonResponse(json_data)

def manage_mon(request):
    user = request.user
    content = {'username':user.tenCanBo}
    if user.is_authenticated and user.position == 0:
        if request.method == 'POST':
            if 'delete' in request.POST:
                Mon.objects.get(id=request.POST['delete']).delete()
            else:
                if request.POST['kieu'] == 'new':
                    try:
                        Mon.objects.create(maMon=request.POST['mamon'], tenMon=request.POST['tenmon'])
                    except:
                        pass
                else:
                    try:
                        m = Mon.objects.get(id=request.POST['id'])
                        m.maMon = request.POST['mamon']
                        m.tenMon = request.POST['tenmon']
                        m.save()
                    except:
                        pass
        return render(request,'adminkt/manager_mon.html',content)
    return redirect('CoiThi:login')

def manage_mon_data(request):
    user = request.user
    if user.is_authenticated and user.position == 0:
        data = []
        for mon in Mon.objects.all():
            ten = '<p id="tenmon_{0}">{1}</p>'.format(mon.id, mon.tenMon)
            ma = '<p id="mamon_{0}">{1}</p>'.format(mon.id, mon.maMon)
            options = '''
                <div class="btn-group">
                    <button type="button" class="btn btn-info" data-toggle="modal" data-target="#new_mon" data-title="edit" id="edit_{0}">
                        <i class="fa fa-cog" data-toggle="tooltip" title="Chỉnh sửa"></i>
                    </button>
                    <button type="button" class="btn btn-danger" data-title="del" id="del_{0}">
                        <i class="fa fa-trash" data-toggle="tooltip" title="Xóa"></i>
                    </button>
                </div>
            '''.format(mon.id)
            data.append([ma,ten, options])
        big_data = {"data": data}
        json_data = json.loads(json.dumps(big_data))
        return JsonResponse(json_data)

def manage_hocvien(request):
    user = request.user
    content = {'username': user.tenCanBo, 'ds_donvi': DonVi.objects.all()}
    if user.is_authenticated and user.position == 0:
        if request.method == 'POST':
            if 'delete' in request.POST:
                SinhVien.objects.get(id=request.POST['delete']).delete()
            elif 'fullname' in request.POST:
                if request.POST['kieu'] == 'new':
                    try:
                        SinhVien.objects.create(tenSinhVien=request.POST['fullname'],
                                                    maSinhVien =request.POST['ma'],
                                                    ngaySinh=request.POST['ngaysinh'],
                                                    maDonVi=DonVi.objects.get(id = request.POST['donvi']))
                            # new_lop = Lop.objects.get(ten=request.POST['list_lop'])
                            # ChiTietLop.objects.create(lop_id=new_lop, myuser_id=hs)
                    except:
                        pass
                else:
                    try:
                        hs = SinhVien.objects.get(id =request.POST['id'])
                        hs.tenSinhVien = request.POST['fullname']
                        hs.maSinhVien = request.POST['ma']
                        hs.ngaySinh = request.POST['ngaysinh']
                        hs.maDonVi = DonVi.objects.get(id = request.POST['donvi'])
                        hs.save()
                    except:
                        pass
            else:
                list_student = request.POST['list_student']
                list_student = json.loads(list_student)
                for stu in list_student:
                    if stu is None:
                        continue
                    try:
                        SinhVien.objects.create(tenSinhVien=stu[0],
                                            maSinhVien =stu[1],
                                            ngaySinh= datetime.datetime.strptime(stu[2], '%Y-%m-%dT%H:%M:%S.%fZ').strftime('%Y-%m-%d'),
                                            maDonVi=DonVi.objects.get(maDonVi = stu[3]))
                    except:
                        continue
        return render(request, 'adminkt/manager_hocvien.html', content)
    else:
        return HttpResponseRedirect('/')

def manage_hocvien_data(request):
    user = request.user
    if user.is_authenticated and user.position == 0:
        data = []
        try:
            ls_student = SinhVien.objects.all()
            now = datetime.datetime.now()
            yearnow = now.year
            for sv in ls_student:
                fullname = '<p id="full_{0}">{1}</p>'.format(sv.id, sv.tenSinhVien)
                masv = '<p id="ma_{}">{}</p>'.format(sv.id,sv.maSinhVien)
                tuoisv = yearnow - sv.ngaySinh.year
                tuoi = '<p id="tuoi_{}">{}</p>'.format(sv.id,tuoisv)
                donvi = '<p id="donvi_{0}">{1}</p>'.format(sv.id,sv.maDonVi.tenDonVi)
                options = '''
                    <div class="btn-group">
                        <button type="button" class="btn btn-info" data-toggle="modal" data-target="#new_student" data-title="edit" id="edit_{0}">
                            <i class="fa fa-cog" data-toggle="tooltip" title="Chỉnh sửa"></i>
                        </button>
                        <button type="button" class="btn btn-danger" data-title="del" id="del_{0}">
                            <i class="fa fa-trash" data-toggle="tooltip" title="Xóa"></i>
                        </button>
                    </div>
                '''.format(sv.id)
                data.append([fullname, masv, tuoi, donvi, options])
        except:
            pass
        big_data = {"data": data}
        json_data = json.loads(json.dumps(big_data))
        return JsonResponse(json_data)
def thongkehocvien_data(request):
    user = request.user
    if user.is_authenticated and user.position == 0:
        data = []
        try:
            ls_student = SinhVien.objects.all()
            now = datetime.datetime.now()
            yearnow = now.year
            for sv in ls_student:
                fullname = '<p id="full_{0}">{1}</p>'.format(sv.id, sv.tenSinhVien)
                masv = '<p id="ma_{}">{}</p>'.format(sv.id,sv.maSinhVien)
                tuoisv = yearnow - sv.ngaySinh.year
                tuoi = '<p id="tuoi_{}">{}</p>'.format(sv.id,tuoisv)
                donvi = '<p id="donvi_{0}">{1}</p>'.format(sv.id,sv.maDonVi.tenDonVi)
                options = '''
                    <div class="btn-group">
                        <button type="button" title="Thống kê điểm" class="thong_ke btn btn-info" data-toggle="modal" data-idsv = {0} data-target="#thongke" data-title="thongke" id="thongke_{0}">
                            <i class="fa fa-line-chart" data-toggle="tooltip" ></i>
                        </button>
                    </div>
                '''.format(sv.id)
                data.append([fullname, masv, tuoi, donvi, options])
        except:
            pass
        big_data = {"data": data}
        json_data = json.loads(json.dumps(big_data))
        return JsonResponse(json_data)
def thongkehocvien_data_diem(request):
    user = request.user
    if user.is_authenticated and user.position == 0:
        data = []
        json_data= []
        try:
            ls_diem = ChiTietLop.objects.filter(maSinhVien = request.POST['idsv'])
            now = datetime.datetime.now()
            i = 1
            for diem in ls_diem:
                stt = i
                mon_hoc =  LopHoc.objects.get(id = diem.maLop.id).maMon.tenMon
                
                monhoc = '<p id="mon">{0}</p>'.format(mon_hoc)
                diem_thi = '<p id="diem">{0}</p>'.format(diem.diem)
                data.append([stt, monhoc, diem_thi])
                i += 1
        except:
            pass
        
        return JsonResponse(data,safe=False)
def log_hocvien_data(request):
    user = request.user
    if user.is_authenticated and user.position == 0:
        data = []
        json_data= []
        try:
            ls_log = ChiTietLop.objects.filter(maLop = LopHoc.objects.get(id=request.POST['lop'])).filter(maSinhVien = SinhVien.objects.get(id=request.POST['hocvien']))
            print(ls_log)
            for log in ls_log:
                for log_diem in log.log_sua_diem.all().order_by('-ngaySua'):
                    fullname = '<p id="full_name">{0}</p>'.format(SinhVien.objects.get(id=request.POST['hocvien']).tenSinhVien)
                    diemCu = '<p id="diemcu">{}</p>'.format(log_diem.diemCu)
                    diemMoi = '<p id="diemmoi">{}</p>'.format(log_diem.diemMoi)
                    ngaysua = '<p id="ngaysua">{}</p>'.format(log_diem.ngaySua)
                    nguoisua = '<p id="nguoisua">{}</p>'.format(log_diem.nguoiSua.tenCanBo)
                    data.append([fullname, diemCu, diemMoi, ngaysua, nguoisua])
        except:
            pass
        return JsonResponse(data,safe=False)

def manager_canbo(request):
    user = request.user
    content = {'username':user.tenCanBo, 'ds_donvi':DonVi.objects.all()}
    if user.is_authenticated and user.position == 0:
        if request.method == 'POST':
            if 'delete' in request.POST:
                CanBo.objects.get(id=request.POST['delete']).delete()
            elif 'block' in request.POST:
                gv = CanBo.objects.get(id=request.POST['block'])
                if gv.is_active:
                    gv.is_active = False
                else:
                    gv.is_active = True
                gv.save()
            elif 'fullname' in request.POST:
                if request.POST['kieu'] == 'new':
                    try:
                        gv = CanBo.objects.create(email=request.POST['email'],
                                            tenCanBo=request.POST['fullname'],
                                            username=request.POST['username'],
                                            quanHam =request.POST['quanham'],
                                            position = 1,
                                            maDonVi = DonVi.objects.get(id = request.POST['donvi']),
                                            maCanBo=request.POST['macanbo'])
                        gv.set_password(request.POST['password'])
                        gv.save()
                    except:
                        pass
                else:
                    try:
                        gv = CanBo.objects.get(username=request.POST['username'])
                        gv.tenCanBo = request.POST['fullname']
                        gv.quanHam = request.POST['quanham']
                        gv.email = request.POST['email']
                        gv.maCanBo = request.POST['macanbo']
                        gv.maDonVi = DonVi.objects.get(id = request.POST['donvi'])
                        gv.save()
                    except:
                        pass
            else:
                list_teacher = request.POST['list_canbo']
                list_teacher = json.loads(list_teacher)
                for i, tea in enumerate(list_teacher):
                    if len(tea) == 0:
                        continue
                    try:
                        gv = CanBo.objects.create(email=tea[2],
                                            tenCanBo=tea[0],
                                            maCanBo=tea[5],
                                            quanHam=tea[4],
                                            username=tea[1],
                                            position = 1,
                                            maDonVi= DonVi.objects.get(id =tea[6]))
                        gv.set_password(tea[3])
                        gv.save()
                    except:
                        continue
        return render(request, 'adminkt/manager_canbo.html', content)
    else:
        return HttpResponseRedirect('/')

def manage_canbo_data(request):
    user = request.user
    if user.is_authenticated and user.position == 0:
        data = []
        try:
            ls_teacher = CanBo.objects.filter()
            for teacher in ls_teacher:
                fullname = '<p id="full_{0}">{1}</p>'.format(teacher.id, teacher.tenCanBo)
                username = '<p id="user_{0}">{1}</p>'.format(teacher.id, teacher.username)
                macanbo = '<p id="macb_{0}">{1}</p>'.format(teacher.id, teacher.maCanBo)
                quanham = '<p id="quanham_{0}">{1}</p>'.format(teacher.id, teacher.quanHam)
                donvi = '<p id="donvi_{0}">{1}</p>'.format(teacher.id,teacher.maDonVi.tenDonVi)
                trang_thai = ''
                if(teacher.is_active):
                    icon = 'fa fa-lock'
                    title = 'khóa'
                    trang_thai = '<span class="label label-success">kích hoạt</span>'
                else:
                    icon = 'fa fa-unlock'
                    title = 'mở khóa'
                    trang_thai = '<span class="label label-danger">khóa</span>'
                options = '''
                    <div class="btn-group">
                        <button type="button" class="btn btn-info" data-toggle="modal" data-target="#new_teacher" data-title="edit" id="edit_{0}">
                            <i class="fa fa-cog" data-toggle="tooltip" title="Chỉnh sửa"></i>
                        </button>
                        <button type="button" class="btn btn-warning" data-title="block" id="block_{0}">
                            <i class="{2}" data-toggle="tooltip" title="{3}"></i></i>
                        </button>
                        <button type="button" class="btn btn-danger" data-title="del" id="del_{0}">
                            <i class="fa fa-trash" data-toggle="tooltip" title="Xóa"></i>
                        </button>
                    </div>
                    <p hidden id="email_{0}">{1}</p>
                '''.format(teacher.id, teacher.email, icon, title)
                data.append([fullname, macanbo, quanham, donvi, username, trang_thai, options])
        except:
            pass
        big_data = {"data": data}
        json_data = json.loads(json.dumps(big_data))
        return JsonResponse(json_data)

def kithi(request):
    if not request.user.is_authenticated:
        return redirect('CoiThi:login')
    return render(request,'adminkt/manager_kithi.html')

def thongke(request):
    user = request.user
    content = {'username':user.tenCanBo,'kithi':KyThi.objects.all()}
    if user.is_authenticated and user.position == 0:
        return render(request, 'adminkt/thongke.html', content)
    else:
        return HttpResponseRedirect('/')
def thongkehocvien(request):
    user = request.user
    content = {'username':user.tenCanBo,'kithi':KyThi.objects.all()}
    if user.is_authenticated and user.position == 0:
        return render(request, 'adminkt/thongkesinhvien.html', content)
    else:
        return HttpResponseRedirect('/')
def thongketheolop(request):
    user = request.user
    content = {'username':user.tenCanBo,'kithi':KyThi.objects.all()}
    if user.is_authenticated and user.position == 0:
        return render(request, 'adminkt/thongketheolop.html', content)
    else:
        return HttpResponseRedirect('/')
def thongke_coi(request):
    user = request.user
    content = {'username':user.tenCanBo,'kithi':KyThi.objects.all()}
    if user.is_authenticated and user.position == 0:
        return render(request, 'adminkt/thongkecoi.html', content)
    else:
        return HttpResponseRedirect('/')
def thongke_thanhtoan(request):
    user = request.user
    content = {'username':user.tenCanBo,'kithi':KyThi.objects.all()}
    if user.is_authenticated and user.position == 0:
        return render(request, 'adminkt/thongkethanhtoancoithi.html', content)
    else:
        return HttpResponseRedirect('/')
def thongke_truot(request):
    user = request.user
    content = {'username':user.tenCanBo,'kithi':KyThi.objects.all()}
    if user.is_authenticated and user.position == 0:
        return render(request, 'adminkt/thongkekithi.html', content)
    else:
        return HttpResponseRedirect('/')
def thongke_vipham(request):
    user = request.user
    content = {'username':user.tenCanBo,'kithi':KyThi.objects.all()}
    if user.is_authenticated and user.position == 0:
        return render(request, 'adminkt/thongkekithi_vipham.html', content)
    else:
        return HttpResponseRedirect('/')
def thongke_hoanthi(request):
    user = request.user
    content = {'username':user.tenCanBo,'kithi':KyThi.objects.all()}
    if user.is_authenticated and user.position == 0:
        return render(request, 'adminkt/thongkekithi_hoanthi.html', content)
    else:
        return HttpResponseRedirect('/')
def thongke_kithi(request,kithi):
    user = request.user
    if user.is_authenticated and user.position == 0:
        data = []
        file = ''
        try:
            ls_kt = KyThi.objects.get(id = kithi)
            ls_pt = PhongThi.objects.filter(maKyThi = ls_kt)
            a = 0
            now = datetime.datetime.now()
            document = Document()
            style = document.styles['Normal']
            font = style.font
            font.name = 'Time New Roman'
            font.size = Pt(14)
            section = document.sections[0]
            section.page_height = Mm(297)
            section.page_width = Mm(210)
            section.left_margin = Mm(15)
            section.right_margin = Mm(15)

            table1 = document.add_table(rows=1, cols=2)
            table1.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table1.rows[0].cells
            paragraph = hdr_cells[0].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run('HỌC VIỆN AN NINH NHÂN DÂN')
            run.bold = True
            run.add_break()
            run1 = paragraph.add_run('PHÒNG KHẢO THÍ & ĐBCLĐT')
            run1.bold = True
            run1.underline = True
            # cell.add_paragraph('')0
            # cell.add_run('HỌC VIỆN AN NINH NHÂN DÂN').bodl
            paragraph1 = hdr_cells[1].paragraphs[0]
            paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = paragraph1.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM')
            run2.bold = True
            run2.add_break()
            run3 = paragraph1.add_run('Độc lập – Tự do – Hạnh phúc')
            run3.bold = True
            run3.underline = True
            text = 'Hà nội, ngày '+ str(now.day) +' tháng ' + str(now.month) + ' năm ' + str(now.year)

            paragraph3 =  document.add_paragraph()
            paragraph3.add_run(text).italic = True
            paragraph3.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            Tieude = 'THỐNG KÊ CHẤM THI KẾT THÚC HỌC PHẦN'

            paragraph3 = document.add_paragraph()
            paragraph3.add_run(Tieude).bold = True
            paragraph3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            Namhoc ='Năm học: 2017 – 2018'

            paragraph3 = document.add_paragraph()
            paragraph3.add_run(Namhoc)
            paragraph3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            body_text_style = document.styles['Table Grid']
            table = document.add_table(rows=2, cols=10,style = body_text_style)
            hdr_cells1 = table.rows[0].cells
            hdr_cells2 = table.rows[1].cells

            hdr_cells1[0].text = 'TT'
            cel_meger0 = hdr_cells1[0].merge(hdr_cells2[0])
            cel_meger0 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[1].text = 'Lớp'
            cel_meger1 = hdr_cells1[1].merge(hdr_cells2[1])
            cel_meger1 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[2].text = 'Quân số'
            cel_meger2 = hdr_cells1[2].merge(hdr_cells2[2])
            cel_meger2 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[3].text = 'Ngày thi'
            cel_meger3 = hdr_cells1[3].merge(hdr_cells2[3])
            cel_meger3 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[4].text = 'Môn thi'
            cel_meger4 = hdr_cells1[4].merge(hdr_cells2[4])
            cel_meger4 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[5].text = 'Hình thức thi'
            cel_meger5 = hdr_cells1[5].merge(hdr_cells2[5])
            cel_meger5 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[6].text = 'Cán bộ chấm thi'
            cel_meger6 = hdr_cells1[6].merge(hdr_cells1[7].merge(hdr_cells1[8]))
            cel_meger6 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells2[6].text = 'Họ và tên'
            hdr_cells2[7].text = 'Đơn vị'
            hdr_cells2[8].text = 'Cấp hàm'

            hdr_cells1[9].text = 'Ghi chú'
            cel_meger7 = hdr_cells1[9].merge(hdr_cells2[9])
            cel_meger7 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            for pt in ls_pt:
                row_cells = table.add_row().cells
                a = a + 1
                row_cells[0].text = str(a)
                lop = '<p id="lop_{0}">{1}</p>'.format(pt.id, pt.maLop.tenLop)
                row_cells[1].text = pt.maLop.tenLop
                quanso = '<p id="quanso_{0}">{1}</p>'.format(pt.id,ChiTietLop.objects.filter(maLop = pt.maLop).count())
                row_cells[2].text = str(ChiTietLop.objects.filter(maLop = pt.maLop).count())
                ngaythi = '<p id="ngaythi_{0}">{1}</p>'.format(pt.id, pt.ngayThi.strftime('%d-%m-%Y'))
                row_cells[3].text = pt.ngayThi.strftime('%d-%m-%Y')
                mon = '<p id="mon_{0}">{1}</p>'.format(pt.id, pt.maLop.maMon.tenMon)
                row_cells[4].text = pt.maLop.maMon.tenMon
                hinhthuc = '<p id="hinhthuc_{0}">{1}</p>'.format(pt.id,pt.hinhThucThi)
                row_cells[5].text = str(pt.hinhThucThi)
                tencb1 = '<p id="tencb1_{0}">{1}</p>'.format(pt.id,ChamThi.objects.get(maPhong = pt.id).canBoCham1.tenCanBo)
                row_cells[6].text = ChamThi.objects.get(maPhong = pt.id).canBoCham1.tenCanBo
                donvi = '<p id="donvi_{0}">{1}</p>'.format(pt.id,ChamThi.objects.get(maPhong = pt.id).canBoCham1.maDonVi.tenDonVi)
                row_cells[7].text = ChamThi.objects.get(maPhong = pt.id).canBoCham1.maDonVi.tenDonVi
                quanham = '<p id="quanham_{0}">{1}</p>'.format(pt.id,ChamThi.objects.get(maPhong = pt.id).canBoCham1.quanHam)
                row_cells[8].text = ChamThi.objects.get(maPhong = pt.id).canBoCham1.quanHam
                ghichu = ''
                row_cells[9].text = ' '
                data.append([str(a),lop, quanso, ngaythi, mon, hinhthuc, tencb1,donvi,quanham,ghichu])

            document.add_paragraph().add_run().add_break()
            table2 = document.add_table(rows=1, cols=2)
            table2.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table2.rows[0].cells
            paragraph6 = hdr_cells[0].paragraphs[0]
            paragraph6.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run5 = paragraph6.add_run('TRƯỞNG PHÒNG')
            run5.bold = True
            paragraph7 = hdr_cells[1].paragraphs[0]
            paragraph7.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run6 = paragraph7.add_run('CÁN BỘ THỐNG KÊ')
            run6.bold = True
            path = 'media/doc/'+Tieude +'.docx'
            file = Tieude +'.docx'
            document.save(path)
        except:
            pass
        big_data = {"data": data,'download' : file}
        json_data = json.loads(json.dumps(big_data))
        return JsonResponse(json_data)
    else:
        return HttpResponseRedirect('/')
def thongke_thanhtoan_coi(request,kithi):
    user = request.user
    if user.is_authenticated and user.position == 0:
        file = ''
        data = []
        tongsobuoi = 0
        try:
            ls_kt = KyThi.objects.get(id = kithi)
            ls_pt = PhongThi.objects.filter(maKyThi = ls_kt)
            a = 0
            now = datetime.datetime.now()
            document = Document()
            style = document.styles['Normal']
            font = style.font
            font.name = 'Time New Roman'
            font.size = Pt(14)
            section = document.sections[0]
            section.page_height = Mm(297)
            section.page_width = Mm(210)
            section.left_margin = Mm(15)
            section.right_margin = Mm(15)
            table1 = document.add_table(rows=1, cols=2)
            table1.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table1.rows[0].cells
            paragraph = hdr_cells[0].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run('HỌC VIỆN AN NINH NHÂN DÂN')
            run.bold = True
            run.add_break()
            run1 = paragraph.add_run('PHÒNG KHẢO THÍ & ĐBCLĐT')
            run1.bold = True
            run1.underline = True
            # cell.add_paragraph('')0
            # cell.add_run('HỌC VIỆN AN NINH NHÂN DÂN').bodl
            paragraph1 = hdr_cells[1].paragraphs[0]
            paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = paragraph1.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM')
            run2.bold = True
            run2.add_break()
            run3 = paragraph1.add_run('Độc lập – Tự do – Hạnh phúc')
            run3.bold = True
            run3.underline = True
            text = 'Hà nội, ngày '+ str(now.day) +' tháng ' + str(now.month) + ' năm ' + str(now.year)
            paragraph3 =  document.add_paragraph()
            paragraph3.add_run(text).italic = True
            paragraph3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            Tieude = 'THỐNG KÊ THANH TOÁN COI THI'
            paragraph3 = document.add_paragraph()
            paragraph3.add_run(Tieude).bold = True
            paragraph3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            Namhoc ='Năm học: 2017 – 2018'
            paragraph3 = document.add_paragraph()
            paragraph3.add_run(Namhoc)
            paragraph3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            body_text_style = document.styles['Table Grid']
            table = document.add_table(rows=1, cols=10,style = body_text_style)
            hdr_cells1 = table.rows[0].cells
            hdr_cells1[0].text = 'TT'
            hdr_cells1[1].text = 'Họ và tên'
            hdr_cells1[2].text = 'Đơn vị'
            hdr_cells1[3].text = 'Cấp hàm'
            hdr_cells1[4].text = 'Số buổi'
            hdr_cells1[5].text = 'Giá biểu'
            hdr_cells1[6].text = 'Thành tiền'
            hdr_cells1[7].text = 'Trừ thuế'
            hdr_cells1[8].text = 'Còn lĩnh'
            hdr_cells1[9].text = 'Kí tên'
            lis =[]
            
            for pt in ls_pt:
                lis.append(pt.canBoCoi1.id)
                lis.append(pt.canBoCoi2.id)
            for x in set(lis):
                row_cells = table.add_row().cells
                a = a + 1
                row_cells[0].text = str(a)
                hoten = '<p id="hoten_{0}">{1}</p>'.format(x,CanBo.objects.get(id = x).tenCanBo)
                row_cells[1].text = CanBo.objects.get(id = x).tenCanBo
                donvi = '<p id="donvi_{0}">{1}</p>'.format(x,CanBo.objects.get(id = x).maDonVi.tenDonVi)
                row_cells[2].text = CanBo.objects.get(id = x).maDonVi.tenDonVi
                capham = '<p id="capham_{0}">{1}</p>'.format(x,CanBo.objects.get(id = x).quanHam)
                row_cells[3].text = CanBo.objects.get(id = x).quanHam
                sobuoi = '<p id="sobuoi_{0}">{1}</p>'.format(x,lis.count(x))
                row_cells[4].text = str(lis.count(x))
                tongsobuoi += int(lis.count(x))
                giabieu = '<p id="giabieu_{0}">{1}</p>'.format(x,'60000')
                row_cells[5].text = str('60000')
                thanhtien = '<p id="thanhtien_{0}">{1}</p>'.format(x,lis.count(x)*60000)
                row_cells[6].text = str(lis.count(x)*60000)
                thue = '<p id="thue_{0}">{1}</p>'.format(x,int((lis.count(x)*60000)*10/100))
                row_cells[7].text = str(int((lis.count(x)*60000)*10/100))
                conlinh = '<p id="conlinh_{0}">{1}</p>'.format(x,int((lis.count(x)*60000)*90/100))
                row_cells[8].text =str(int((lis.count(x)*60000)*90/100))
                kinhan = ''
                row_cells[9].text = ' '
                data.append([str(a),hoten, donvi, capham, sobuoi,giabieu,thanhtien,thue,conlinh,kinhan])
            row_cells = table.add_row().cells
            row_cells[0].text = '#'
            row_cells[1].text = 'Tổng'
            cel_meger8 = row_cells[1].merge(row_cells[2].merge(row_cells[3]))
            cel_meger8 = '\n'.join(
                cell.text for cell in row_cells if cell.text
            )
            row_cells[4].text = str(tongsobuoi)
            row_cells[6].text = str(int(tongsobuoi*60000))
            row_cells[7].text = str(int(tongsobuoi*60000*0.1))
            row_cells[8].text = str(int(tongsobuoi*60000*0.9))
            document.add_paragraph().add_run().add_break()
            table2 = document.add_table(rows=1, cols=2)
            table2.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table2.rows[0].cells
            paragraph6 = hdr_cells[0].paragraphs[0]
            paragraph6.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run5 = paragraph6.add_run('TRƯỞNG PHÒNG')
            run5.bold = True
            paragraph7 = hdr_cells[1].paragraphs[0]
            paragraph7.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run6 = paragraph7.add_run('CÁN BỘ THỐNG KÊ')
            run6.bold = True
            path = 'media/doc/'+Tieude +'.docx'
            file = Tieude +'.docx'
            document.save(path)
        except:
            pass
        big_data = {"data": data,'download' : file,"sobuoi" : tongsobuoi,"tongtien":str(tongsobuoi*60000),"tongthue":str(int(tongsobuoi*60000*0.1)),"tonglinh" : str(int(tongsobuoi*60000*0.9))}
        json_data = json.loads(json.dumps(big_data))
        return JsonResponse(json_data)
    else:
        return HttpResponseRedirect('/')

def manager_kithi(request):
    user = request.user
    content = {'username': user.tenCanBo,'ds_cb':CanBo.objects.all(),'ds_lop':LopHoc.objects.all(),'ds_kithi':KyThi.objects.all()}
    if user.is_authenticated and user.position == 0:
        if request.method == 'POST':
            if 'delete' in request.POST:
                KyThi.objects.get(id=request.POST['delete']).delete()
            else:
                if request.POST['kieu'] == 'new':
                    try:
                        KyThi.objects.create(tenKyThi=request.POST['tenkithi'], ngayBatDau=request.POST['tungay'], ngayKetThuc=request.POST['denngay'])
                    except:
                        pass
                else:
                    try:
                        m = KyThi.objects.get(id=request.POST['id'])
                        m.tenKyThi = request.POST['tenkithi']
                        m.ngayBatDau = request.POST['tungay']
                        m.ngayKetThuc = request.POST['denngay']
                        m.save()
                    except:
                        pass
        return render(request, 'adminkt/manager_kithi.html', content)
    else:
        return HttpResponseRedirect('/')

def manage_kithi_data(request):
    user = request.user
    if user.is_authenticated and user.position == 0:
        data = []
        try:
            ls_kithi = KyThi.objects.all()
            
            for kt in ls_kithi:
                tenkithi = '<p id="tenkithi_{0}">{1}</p>'.format(kt.id, kt.tenKyThi)
                tungay = '<p id="tungay_{}" data-tungay="{}">{}</p>'.format(kt.id,kt.ngayBatDau,kt.ngayBatDau.strftime('%d-%m-%Y'))
                denngay = '<p id="denngay_{0}" data-denngay="{1}">{2}</p>'.format(kt.id,kt.ngayKetThuc,kt.ngayKetThuc.strftime('%d-%m-%Y'))
                ds_phongthi = '''
                {1}  <i class="fa fa-info-circle" data-title="{0}" data-toggle="modal" data-target="#detail_room"></i>
                '''.format(kt.tenKyThi,PhongThi.objects.filter(maKyThi_id = kt.id).count())
                options = '''
                    <div class="btn-group">
                        <button type="button" class="btn btn-info" data-toggle="modal" data-target="#new_kithi" data-title="edit" id="edit_{0}">
                            <i class="fa fa-cog" data-toggle="tooltip" title="Chỉnh sửa"></i>   Chỉnh sửa</button>
                        <button type="button" class="btn btn-danger" data-toggle="modal" data-title="del" id="del_{0}" data-target="#delete_kithi">
                            <i class="fa fa-trash" data-toggle="tooltip" title="Xóa"></i>
                        </button>
                    </div>
                '''.format(kt.id)
                data.append([tenkithi, tungay, denngay, ds_phongthi, options])
        except:
            pass
        big_data = {"data": data}
        json_data = json.loads(json.dumps(big_data))
        return JsonResponse(json_data)

def manage_ki_thi_data(request,kithi):
    user = request.user
    if user.is_authenticated and user.position == 0:
        data = []
        try:
            ls_chi_tiet = KyThi.objects.filter(tenKyThi = kithi).values('id')
            ls_phongthi = PhongThi.objects.filter(maKyThi__in = ls_chi_tiet)
            
            for room in ls_phongthi:
                tenphong = '<p id="room_{0}">{1}</p>'.format(room.id, room.tenPhong)
                vitri = '<p id="vitri_{0}">{1}</p>'.format(room.id, room.viTri)
                ngaythi = '<p id="ngaythi_{0}">{1}</p>'.format(room.id, room.ngayThi.strftime('%d-%m-%y'))
                gio = '<p id="gio_{0}">{1}</p>'.format(room.id, room.gio)
                tenlop = '<p id="tenlop_{0}">{1}</p>'.format(room.id, room.maLop.tenLop)
                canbocoi1 = '<p id="cbcoi1_{0}">{1}</p>'.format(room.id, room.canBoCoi1.tenCanBo)
                canBoCoi2 = '<p id="cbcoi2_{0}">{1}</p>'.format(room.id, room.canBoCoi2.tenCanBo)
                option = '''
                    <div class="btn-group">
                        <button type="button" class="btn btn-danger" data-title="del" id="del_{0}">
                            <i class="fa fa-trash" data-toggle="tooltip" title="Xóa"></i>
                        </button>
                    </div>
                '''.format(room.id)
                data.append([tenphong, vitri, ngaythi,gio,tenlop,canbocoi1,canBoCoi2,option])
        except:
            pass
        big_data = {"data": data}
        json_data = json.loads(json.dumps(big_data))
        return JsonResponse(json_data)

def manage_add_kithi(request):
    user = request.user
    if user.is_authenticated and user.position == 0:
        if request.method == 'POST':
            if 'delete' in request.POST:
                PhongThi.objects.get(id=request.POST['delete']).delete()
            elif 'lop' in request.POST:

                if request.POST['kieu'] == 'new':
                    try:
                        PhongThi.objects.create(maKyThi = KyThi.objects.get(id = request.POST['kithi']) ,
                                                canBoCoi1=CanBo.objects.get(id =request.POST['cbc1']),
                                                canBoCoi2=CanBo.objects.get(id =request.POST['cbc2']),
                                                maLop = LopHoc.objects.get(id = request.POST['lop']),
                                                tenPhong = request.POST['tenphong'],
                                                viTri = request.POST['vitri'],
                                                ngayThi = request.POST['ngaythi'],
                                                gio = request.POST['gio'],
                                                hinhThucThi = request.POST['hinhthuc'],
                                                thoiGianThi = request.POST['thoigian'])

                        ChamThi.objects.create(maPhong = PhongThi.objects.get(tenPhong = request.POST['tenphong']),
                                            canBoCham1=CanBo.objects.get(id =request.POST['cbcham1']),
                                           canBoCham2=CanBo.objects.get(id =request.POST['cbcham2']) )
                    except:
                        pass
        return JsonResponse({'status':'succsec'})
    else:
        return HttpResponseRedirect('/')

def download(request, pathdoc):
    try:
        file_path = "media/doc/" + pathdoc
        response = HttpResponse(open(file_path, 'rb').read())
        response['Content-Type'] = 'application/vnd.ms-word'
        response['Content-Disposition'] = 'attachment; filename='+ pathdoc
        
    except:
        response= ''
    return response

def thongke_kithi_coi(request,kithi):
    user = request.user
    if user.is_authenticated and user.position == 0:
        data = []
        file = ''
        a = 0
        b = 0
        try:
            ls_kt = KyThi.objects.get(id = kithi)
            ls_pt = PhongThi.objects.filter(maKyThi = ls_kt)
            document = Document()
            style = document.styles['Normal']
            font = style.font
            font.name = 'Time New Roman'
            font.size = Pt(14)
            section = document.sections[0]
            section.page_height = Mm(297)
            section.page_width = Mm(210)
            section.left_margin = Mm(15)
            section.right_margin = Mm(15)

            table1 = document.add_table(rows=1, cols=2)
            table1.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table1.rows[0].cells
            paragraph = hdr_cells[0].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run('HỌC VIỆN AN NINH NHÂN DÂN')
            run.bold = True
            run.add_break()
            run1 = paragraph.add_run('PHÒNG KHẢO THÍ & ĐBCLĐT')
            run1.bold = True
            run1.underline = True
            paragraph1 = hdr_cells[1].paragraphs[0]
            paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = paragraph1.add_run('THỐNG KÊ CÁN BỘ COI THI HỌC PHẦN')
            run2.bold = True
            run2.add_break()
            run3 = paragraph1.add_run('HỌC KỲ……, NĂM HỌC 20…. – 20….')
            run3.bold = True
            run3.underline = True
            
            Tieude = 'THỐNG KÊ COI THI KẾT THÚC HỌC PHẦN'
            body_text_style = document.styles['Table Grid']
            table = document.add_table(rows=2, cols=10,style = body_text_style)
            hdr_cells1 = table.rows[0].cells
            hdr_cells2 = table.rows[1].cells

            hdr_cells1[0].text = 'TT'
            cel_meger0 = hdr_cells1[0].merge(hdr_cells2[0])
            cel_meger0 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[1].text = 'Lớp'
            cel_meger1 = hdr_cells1[1].merge(hdr_cells2[1])
            cel_meger1 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[2].text = 'Quân số'
            cel_meger2 = hdr_cells1[2].merge(hdr_cells2[2])
            cel_meger2 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[3].text = 'Ngày thi'
            cel_meger3 = hdr_cells1[3].merge(hdr_cells2[3])
            cel_meger3 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[4].text = 'Môn thi'
            cel_meger4 = hdr_cells1[4].merge(hdr_cells2[4])
            cel_meger4 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[5].text = 'Hình thức thi'
            cel_meger5 = hdr_cells1[5].merge(hdr_cells2[5])
            cel_meger5 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[6].text = 'Cán bộ chấm thi'
            cel_meger6 = hdr_cells1[6].merge(hdr_cells1[7].merge(hdr_cells1[8]))
            cel_meger6 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells2[6].text = 'Họ và tên'
            hdr_cells2[7].text = 'Đơn vị'
            hdr_cells2[8].text = 'Cấp hàm'

            hdr_cells1[9].text = 'Ghi chú'
            cel_meger7 = hdr_cells1[9].merge(hdr_cells2[9])
            cel_meger7 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            
            for pt in ls_pt:
                row_cells = table.add_row().cells
                a = a + 1
                b += ChiTietLop.objects.all().filter(maLop = pt.maLop).exclude(trangThai ='Hoãn thi').count()
                lop = '<p id="lop_{0}">{1}</p>'.format(pt.id, pt.maLop.tenLop)
                row_cells[1].text = pt.maLop.tenLop
                quanso = '<p id="quanso_{0}">{1}</p>'.format(pt.id,ChiTietLop.objects.filter(maLop = pt.maLop).count())
                row_cells[2].text = str(ChiTietLop.objects.filter(maLop = pt.maLop).count())
                ngaythi = '<p id="ngaythi_{0}">{1}</p>'.format(pt.id, pt.ngayThi.strftime('%d-%m-%Y'))
                row_cells[3].text = pt.ngayThi.strftime('%d-%m-%Y')
                mon = '<p id="mon_{0}">{1}</p>'.format(pt.id, pt.maLop.maMon.tenMon)
                row_cells[4].text = pt.maLop.maMon.tenMon
                hinhthuc = '<p id="hinhthuc_{0}">{1}</p>'.format(pt.id,pt.hinhThucThi)
                row_cells[5].text = str(pt.hinhThucThi)
                tencb1 = '<p id="tencb1_{0}">{1}</p>'.format(pt.id,pt.canBoCoi1.tenCanBo)
                tencb2 = '<p id="tencb2_{0}">{1}</p>'.format(pt.id,pt.canBoCoi2.tenCanBo)
                row_cells[6].text = pt.canBoCoi1.tenCanBo
                donvi1 = '<p id="donvi_{0}">{1}</p>'.format(pt.id,pt.canBoCoi1.maDonVi.tenDonVi)
                donvi2 = '<p id="donvi_{0}">{1}</p>'.format(pt.id,pt.canBoCoi2.maDonVi.tenDonVi)
                row_cells[7].text = pt.canBoCoi1.maDonVi.tenDonVi
                quanham1 = '<p id="quanham_{0}">{1}</p>'.format(pt.id,pt.canBoCoi1.quanHam)
                quanham2 = '<p id="quanham_{0}">{1}</p>'.format(pt.id,pt.canBoCoi2.quanHam)
                row_cells[8].text = pt.canBoCoi1.quanHam
                ghichu = ''
                row_cells[9].text = ' '
                data.append([str(a),lop, quanso, ngaythi, mon, hinhthuc, tencb1,donvi1,quanham1,tencb2,donvi2,quanham2,ghichu])

            document.add_paragraph().add_run().add_break()
            table2 = document.add_table(rows=1, cols=2)
            table2.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table2.rows[0].cells
            paragraph6 = hdr_cells[0].paragraphs[0]
            paragraph6.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run5 = paragraph6.add_run('TRƯỞNG PHÒNG')
            run5.bold = True
            paragraph7 = hdr_cells[1].paragraphs[0]
            paragraph7.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run6 = paragraph7.add_run('CÁN BỘ THỐNG KÊ')
            run6.bold = True
            path = 'media/doc/'+Tieude +'.docx'
            file = Tieude +'.docx'
            document.save(path)
        except:
            pass
        big_data = {"data": data,'download' : file,'solop': a ,'sohocvien': b}
        json_data = json.loads(json.dumps(big_data))
        return JsonResponse(json_data)
    else:
        return HttpResponseRedirect('/')

def thongke_kithi_truot(request,kithi):
    user = request.user
    if user.is_authenticated and user.position == 0:
        data = []
        a = 0
        file = ''
        try:
            ls_kt = KyThi.objects.get(id = kithi)
            ls_pt = PhongThi.objects.filter(maKyThi = ls_kt) 
            
            document = Document()
            style = document.styles['Normal']
            font = style.font
            font.name = 'Time New Roman'
            font.size = Pt(14)
            section = document.sections[0]
            section.page_height = Mm(297)
            section.page_width = Mm(210)
            section.left_margin = Mm(15)
            section.right_margin = Mm(15)

            table1 = document.add_table(rows=1, cols=2)
            table1.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table1.rows[0].cells
            paragraph = hdr_cells[0].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run('HỌC VIỆN AN NINH NHÂN DÂN')
            run.bold = True
            run.add_break()
            run1 = paragraph.add_run('PHÒNG KHẢO THÍ & ĐBCLĐT')
            run1.bold = True
            run1.underline = True
            # cell.add_paragraph('')0
            # cell.add_run('HỌC VIỆN AN NINH NHÂN DÂN').bodl
            paragraph1 = hdr_cells[1].paragraphs[0]
            paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = paragraph1.add_run('THỐNG KÊ SỐ LIỆU HỌC VIÊN TRƯỢT')
            run2.bold = True
            run2.add_break()
            run3 = paragraph1.add_run('HỌC KỲ……, NĂM HỌC 20…. – 20….')
            run3.bold = True
            run3.underline = True
            Tieude = 'THỐNG KÊ SỐ LIỆU HỌC VIÊN TRƯỢT'
            body_text_style = document.styles['Table Grid']
            table = document.add_table(rows=2, cols=8,style = body_text_style)
            hdr_cells1 = table.rows[0].cells
            hdr_cells2 = table.rows[1].cells

            hdr_cells1[0].text = 'TT'
            cel_meger0 = hdr_cells1[0].merge(hdr_cells2[0])
            cel_meger0 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[1].text = 'Lớp'
            cel_meger1 = hdr_cells1[1].merge(hdr_cells2[1])
            cel_meger1 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[2].text = 'Ngày thi'
            cel_meger2 = hdr_cells1[2].merge(hdr_cells2[2])
            cel_meger2 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[3].text = 'Môn thi'
            cel_meger3 = hdr_cells1[3].merge(hdr_cells2[3])
            cel_meger3 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[4].text = 'Hình thức thi'
            cel_meger4 = hdr_cells1[4].merge(hdr_cells2[4])
            cel_meger4 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[5].text = 'Học viên trượt'
            cel_meger6 = hdr_cells1[5].merge(hdr_cells1[6])
            cel_meger6 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells2[5].text = 'Họ và tên'
            hdr_cells2[6].text = 'Điểm'

            hdr_cells1[7].text = 'Ghi chú'
            cel_meger7 = hdr_cells1[7].merge(hdr_cells2[7])
            cel_meger7 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            b = 0
            for phongthi in ls_pt:
                ls_sv = ChiTietLop.objects.filter(diem__lt = 5,maLop = phongthi.maLop.id)
                for pt in ls_sv:
                    row_cells = table.add_row().cells
                    a = a + 1
                    row_cells[0].text = str(a)
                    lop = '<p id="lop_{0}">{1}</p>'.format(pt.id, pt.maLop.tenLop)
                    row_cells[1].text = pt.maLop.tenLop
                    ngaythi = '<p id="ngaythi_{0}">{1}</p>'.format(pt.id,PhongThi.objects.get(maLop = pt.maLop.id).ngayThi.strftime('%d-%m-%Y'))
                    row_cells[2].text = PhongThi.objects.get(maLop = pt.maLop.id).ngayThi.strftime('%d-%m-%Y')
                    mon = '<p id="mon_{0}">{1}</p>'.format(pt.id, pt.maLop.maMon.tenMon)
                    row_cells[3].text = pt.maLop.maMon.tenMon
                    hinhthuc = '<p id="hinhthuc_{0}">{1}</p>'.format(pt.id,PhongThi.objects.get(maLop = pt.maLop.id).hinhThucThi)
                    row_cells[4].text = PhongThi.objects.get(maLop = pt.maLop.id).hinhThucThi
                    tensv = '<p id="tensv_{0}">{1}</p>'.format(pt.id,pt.maSinhVien.tenSinhVien)
                    row_cells[5].text = pt.maSinhVien.tenSinhVien
                    diemthi = '<p id="diem_{0}">{1}</p>'.format(pt.id,pt.diem)
                    # cells[6].text = pt.diem
                    ghichu = '<p id="lydo_{0}">{1}</p>'.format(pt.id,pt.ghiChu)
                    row_cells[7].text = ''
                    data.append([str(a),lop, ngaythi, mon, hinhthuc,tensv,diemthi,ghichu])

            document.add_paragraph().add_run().add_break()
            table2 = document.add_table(rows=1, cols=2)
            table2.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table2.rows[0].cells
            paragraph6 = hdr_cells[0].paragraphs[0]
            paragraph6.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run5 = paragraph6.add_run('TRƯỞNG PHÒNG')
            run5.bold = True
            paragraph7 = hdr_cells[1].paragraphs[0]
            paragraph7.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run6 = paragraph7.add_run('CÁN BỘ THỐNG KÊ')
            run6.bold = True
            path = 'media/doc/'+Tieude +'.docx'
            file = Tieude +'.docx'
            document.save(path)
        except:
            pass
        big_data = {"data": data,'download' : file,'so': a }
        json_data = json.loads(json.dumps(big_data))
        return JsonResponse(json_data)
    else:
        return HttpResponseRedirect('/')

def thongke_kithi_vipham(request,kithi):
    user = request.user
    if user.is_authenticated and user.position == 0:
        data = []
        a = 0
        file = ''
        try:
            ls_kt = KyThi.objects.get(id = kithi)
            ls_pt = PhongThi.objects.filter(maKyThi = ls_kt)
            # ls_sv = ChiTietLop.objects.all().filter(trangThai= 'Vi phạm')
            
            document = Document()
            style = document.styles['Normal']
            font = style.font
            font.name = 'Time New Roman'
            font.size = Pt(14)
            section = document.sections[0]
            section.page_height = Mm(297)
            section.page_width = Mm(210)
            section.left_margin = Mm(15)
            section.right_margin = Mm(15)

            table1 = document.add_table(rows=1, cols=2)
            table1.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table1.rows[0].cells
            paragraph = hdr_cells[0].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run('HỌC VIỆN AN NINH NHÂN DÂN')
            run.bold = True
            run.add_break()
            run1 = paragraph.add_run('PHÒNG KHẢO THÍ & ĐBCLĐT')
            run1.bold = True
            run1.underline = True
            # cell.add_paragraph('')0
            # cell.add_run('HỌC VIỆN AN NINH NHÂN DÂN').bodl
            paragraph1 = hdr_cells[1].paragraphs[0]
            paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = paragraph1.add_run('THỐNG KÊ SỐ LIỆU HỌC VIÊN VI PHẠM')
            run2.bold = True
            run2.add_break()
            run3 = paragraph1.add_run('HỌC KỲ……, NĂM HỌC 20…. – 20….')
            run3.bold = True
            run3.underline = True
            Tieude = 'THỐNG KÊ SỐ LIỆU HỌC VIÊN VI PHẠM'
            body_text_style = document.styles['Table Grid']
            table = document.add_table(rows=2, cols=8,style = body_text_style)
            hdr_cells1 = table.rows[0].cells
            hdr_cells2 = table.rows[1].cells

            hdr_cells1[0].text = 'TT'
            cel_meger0 = hdr_cells1[0].merge(hdr_cells2[0])
            cel_meger0 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[1].text = 'Lớp'
            cel_meger1 = hdr_cells1[1].merge(hdr_cells2[1])
            cel_meger1 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[2].text = 'Ngày thi'
            cel_meger2 = hdr_cells1[2].merge(hdr_cells2[2])
            cel_meger2 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[3].text = 'Môn thi'
            cel_meger3 = hdr_cells1[3].merge(hdr_cells2[3])
            cel_meger3 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[4].text = 'Hình thức thi'
            cel_meger4 = hdr_cells1[4].merge(hdr_cells2[4])
            cel_meger4 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[5].text = 'Học viên vi phạm'
            cel_meger6 = hdr_cells1[5].merge(hdr_cells1[6])
            cel_meger6 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells2[5].text = 'Họ và tên'
            hdr_cells2[6].text = 'Lý do'

            hdr_cells1[7].text = 'Ghi chú'
            cel_meger7 = hdr_cells1[7].merge(hdr_cells2[7])
            cel_meger7 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            b = 0
            for phongthi in ls_pt:
                ls_sv = ChiTietLop.objects.filter(trangThai= 'Vi phạm',maLop = phongthi.maLop.id)
                for pt in ls_sv:
                    row_cells = table.add_row().cells
                    a = a + 1
                    row_cells[0].text = str(a)
                    lop = '<p id="lop_{0}">{1}</p>'.format(pt.id, pt.maLop.tenLop)
                    row_cells[1].text = pt.maLop.tenLop
                    ngaythi = '<p id="ngaythi_{0}">{1}</p>'.format(pt.id,PhongThi.objects.get(maLop = pt.maLop).ngayThi.strftime('%d-%m-%Y'))
                    row_cells[2].text = PhongThi.objects.get(maLop = pt.maLop).ngayThi.strftime('%d-%m-%Y')
                    mon = '<p id="mon_{0}">{1}</p>'.format(pt.id, pt.maLop.maMon.tenMon)
                    row_cells[3].text = pt.maLop.maMon.tenMon
                    hinhthuc = '<p id="hinhthuc_{0}">{1}</p>'.format(pt.id,PhongThi.objects.get(maLop = pt.maLop).hinhThucThi)
                    row_cells[4].text = PhongThi.objects.get(maLop = pt.maLop).hinhThucThi
                    tensv = '<p id="tensv_{0}">{1}</p>'.format(pt.id,pt.maSinhVien.tenSinhVien)
                    row_cells[5].text = pt.maSinhVien.tenSinhVien
                    lydo = '<p id="lydo_{0}">{1}</p>'.format(pt.id,pt.lyDo)
                    row_cells[6].text = pt.lyDo
                    ghichu = '<p id="lydo_{0}">{1}</p>'.format(pt.id,pt.ghiChu)
                    row_cells[7].text = pt.ghiChu
                    data.append([str(a),lop, ngaythi, mon, hinhthuc,tensv,lydo,ghichu])

            document.add_paragraph().add_run().add_break()
            table2 = document.add_table(rows=1, cols=2)
            table2.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table2.rows[0].cells
            paragraph6 = hdr_cells[0].paragraphs[0]
            paragraph6.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run5 = paragraph6.add_run('TRƯỞNG PHÒNG')
            run5.bold = True
            paragraph7 = hdr_cells[1].paragraphs[0]
            paragraph7.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run6 = paragraph7.add_run('CÁN BỘ THỐNG KÊ')
            run6.bold = True
            path = 'media/doc/'+Tieude +'.docx'
            file = Tieude +'.docx'
            document.save(path)
        except:
            pass
        big_data = {"data": data,'download' : file,'so': a }
        json_data = json.loads(json.dumps(big_data))
        return JsonResponse(json_data)
    else:
        return HttpResponseRedirect('/')

def thongke_kithi_hoanthi(request,kithi):
    user = request.user
    if user.is_authenticated and user.position == 0:
        data = []
        a = 0
        file =''
        try:
            ls_kt = KyThi.objects.get(id = kithi)
            ls_pt = PhongThi.objects.filter(maKyThi = ls_kt)
            
            document = Document()
            style = document.styles['Normal']
            font = style.font
            font.name = 'Time New Roman'
            font.size = Pt(14)
            section = document.sections[0]
            section.page_height = Mm(297)
            section.page_width = Mm(210)
            section.left_margin = Mm(15)
            section.right_margin = Mm(15)

            table1 = document.add_table(rows=1, cols=2)
            table1.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table1.rows[0].cells
            paragraph = hdr_cells[0].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run('HỌC VIỆN AN NINH NHÂN DÂN')
            run.bold = True
            run.add_break()
            run1 = paragraph.add_run('PHÒNG KHẢO THÍ & ĐBCLĐT')
            run1.bold = True
            run1.underline = True
            # cell.add_paragraph('')0
            # cell.add_run('HỌC VIỆN AN NINH NHÂN DÂN').bodl
            paragraph1 = hdr_cells[1].paragraphs[0]
            paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = paragraph1.add_run('THỐNG KÊ SỐ LIỆU HỌC VIÊN HOÃN THI')
            run2.bold = True
            run2.add_break()
            run3 = paragraph1.add_run('HỌC KỲ……, NĂM HỌC 20…. – 20….')
            run3.bold = True
            run3.underline = True
            Tieude = 'THỐNG KÊ SỐ LIỆU HỌC VIÊN HÕAN THI'
            body_text_style = document.styles['Table Grid']
            table = document.add_table(rows=2, cols=8,style = body_text_style)
            hdr_cells1 = table.rows[0].cells
            hdr_cells2 = table.rows[1].cells

            hdr_cells1[0].text = 'TT'
            cel_meger0 = hdr_cells1[0].merge(hdr_cells2[0])
            cel_meger0 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[1].text = 'Lớp'
            cel_meger1 = hdr_cells1[1].merge(hdr_cells2[1])
            cel_meger1 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[2].text = 'Ngày thi'
            cel_meger2 = hdr_cells1[2].merge(hdr_cells2[2])
            cel_meger2 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[3].text = 'Môn thi'
            cel_meger3 = hdr_cells1[3].merge(hdr_cells2[3])
            cel_meger3 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[4].text = 'Hình thức thi'
            cel_meger4 = hdr_cells1[4].merge(hdr_cells2[4])
            cel_meger4 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells1[5].text = 'Học viên hoãn thi'
            cel_meger6 = hdr_cells1[5].merge(hdr_cells1[6])
            cel_meger6 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            hdr_cells2[5].text = 'Họ và tên'
            hdr_cells2[6].text = 'lý do'

            hdr_cells1[7].text = 'Ghi chú'
            cel_meger7 = hdr_cells1[7].merge(hdr_cells2[7])
            cel_meger7 = '\n'.join(
                cell.text for cell in hdr_cells1 if cell.text
            )
            b = 0
            for phongthi in ls_pt:
                ls_sv = ChiTietLop.objects.filter(trangThai= 'Hoãn thi',maLop = phongthi.maLop.id)
                for pt in ls_sv:
                    row_cells = table.add_row().cells
                    a = a + 1
                    row_cells[0].text = str(a)
                    lop = '<p id="lop_{0}">{1}</p>'.format(pt.id, pt.maLop.tenLop)
                    row_cells[1].text = pt.maLop.tenLop
                    ngaythi = '<p id="ngaythi_{0}">{1}</p>'.format(pt.id,PhongThi.objects.get(maLop = pt.maLop).ngayThi.strftime('%d-%m-%Y'))
                    row_cells[2].text = PhongThi.objects.get(maLop = pt.maLop).ngayThi.strftime('%d-%m-%Y')
                    mon = '<p id="mon_{0}">{1}</p>'.format(pt.id, pt.maLop.maMon.tenMon)
                    row_cells[3].text = pt.maLop.maMon.tenMon
                    hinhthuc = '<p id="hinhthuc_{0}">{1}</p>'.format(pt.id,PhongThi.objects.get(maLop = pt.maLop).hinhThucThi)
                    row_cells[4].text = PhongThi.objects.get(maLop = pt.maLop).hinhThucThi
                    tensv = '<p id="tensv_{0}">{1}</p>'.format(pt.id,pt.maSinhVien.tenSinhVien)
                    row_cells[5].text = pt.maSinhVien.tenSinhVien
                    lydo = '<p id="lydo_{0}">{1}</p>'.format(pt.id,pt.lyDo)
                    row_cells[6].text = pt.lyDo
                    ghichu = '<p id="lydo_{0}">{1}</p>'.format(pt.id,pt.ghiChu)
                    row_cells[7].text = pt.ghiChu
                    data.append([str(a),lop, ngaythi, mon, hinhthuc,tensv,lydo,ghichu])

            document.add_paragraph().add_run().add_break()
            table2 = document.add_table(rows=1, cols=2)
            table2.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table2.rows[0].cells
            paragraph6 = hdr_cells[0].paragraphs[0]
            paragraph6.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run5 = paragraph6.add_run('TRƯỞNG PHÒNG')
            run5.bold = True
            paragraph7 = hdr_cells[1].paragraphs[0]
            paragraph7.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run6 = paragraph7.add_run('CÁN BỘ THỐNG KÊ')
            run6.bold = True
            path = 'media/doc/'+Tieude +'.docx'
            file = Tieude +'.docx'
            document.save(path)
        except:
            pass
        big_data = {"data": data,'download' : file,'so': a }
        json_data = json.loads(json.dumps(big_data))
        return JsonResponse(json_data)
    else:
        return HttpResponseRedirect('/')
